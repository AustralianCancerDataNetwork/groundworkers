"""DOCX decomposer for tables and heading-grouped text fallback."""

from __future__ import annotations

import io
import re

from groundworkers.services.source_planning.models import RawTable, SourceFormat

_MAX_SAMPLE = 5
_W_VMERGE = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}vMerge"
_W_TCPR = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcPr"
_W_VAL = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val"
_MULTI_SPACE = re.compile(r"\s+")


def decompose(content: bytes, filename: str | None = None) -> list[RawTable]:
    """Parse DOCX bytes and return one ``RawTable`` per Word table."""

    from docx import Document

    try:
        doc = Document(io.BytesIO(content))
    except Exception:
        return []

    results: list[RawTable] = []
    for index, table in enumerate(doc.tables):
        rows_raw = _read_table(table)
        if len(rows_raw) < 2:
            continue
        headers = _normalise_headers(rows_raw[0])
        if not headers or all(header == "" for header in headers):
            continue
        rows = [
            {headers[j]: (row[j] if j < len(row) else "") for j in range(len(headers))}
            for row in rows_raw[1:]
            if any(row)
        ]
        if not rows:
            continue
        results.append(
            RawTable(
                name=_table_name(filename, index),
                headers=headers,
                rows=rows,
                sample_rows=rows[:_MAX_SAMPLE],
                source_format=SourceFormat.DOCX,
                row_count=len(rows),
                metadata={"docx_table_index": index},
            )
        )

    if results:
        return results

    text_rows = _extract_paragraph_sections(doc)
    if text_rows:
        results.append(
            RawTable(
                name=_table_name(filename, 0),
                headers=["section", "text"],
                rows=text_rows,
                sample_rows=text_rows[:_MAX_SAMPLE],
                source_format=SourceFormat.DOCX,
                row_count=len(text_rows),
                metadata={"docx_source": "paragraphs"},
            )
        )
    return results


def _cell(value: str) -> str:
    return _MULTI_SPACE.sub(" ", (value or "")).strip()


def _is_vmerge_continuation(cell) -> bool:
    tc_pr = cell._tc.find(_W_TCPR)
    if tc_pr is None:
        return False
    vmerge = tc_pr.find(_W_VMERGE)
    if vmerge is None:
        return False
    return vmerge.get(_W_VAL) != "restart"


def _read_table(table) -> list[list[str]]:
    carry: dict[int, str] = {}
    rows: list[list[str]] = []
    for row in table.rows:
        row_values: list[str] = []
        for col_idx, cell in enumerate(row.cells):
            if _is_vmerge_continuation(cell):
                row_values.append(carry.get(col_idx, ""))
            else:
                text = _cell(cell.text)
                carry[col_idx] = text
                row_values.append(text)
        rows.append(row_values)
    return rows


def _normalise_headers(raw: list[str]) -> list[str]:
    seen: dict[str, int] = {}
    out: list[str] = []
    for header in raw:
        header = _cell(header) or "col"
        if header in seen:
            seen[header] += 1
            header = f"{header}_{seen[header]}"
        else:
            seen[header] = 1
        out.append(header)
    return out


def _extract_paragraph_sections(doc: object) -> list[dict[str, str]]:
    rows: list[dict[str, str]] = []
    current_heading = ""
    buffer: list[str] = []

    def _flush() -> None:
        body = " ".join(buffer).strip()
        if body:
            rows.append({"section": current_heading, "text": body})
        buffer.clear()

    for para in doc.paragraphs:  # type: ignore[attr-defined]
        text = _cell(para.text)
        if not text:
            continue
        style = (para.style.name or "").lower()
        if "heading" in style:
            _flush()
            current_heading = text
        else:
            buffer.append(text)

    _flush()
    return rows


def _table_name(filename: str | None, index: int) -> str:
    base = ""
    if filename:
        base = filename.rsplit("/", 1)[-1].rsplit("\\", 1)[-1]
        if "." in base:
            base = base.rsplit(".", 1)[0]
    return f"{base}_table{index + 1}" if base else f"table{index + 1}"
